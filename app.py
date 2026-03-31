"""
Katalyst Judgment Radar v5 Final
Google-first approach for all courts. Liberal classification. Wide searches.
"""

import os
import json
import sqlite3
import hashlib
import logging
import time
import re
import threading
from datetime import datetime, timedelta
from urllib.parse import quote, urljoin

import requests
import feedparser
from flask import Flask, jsonify, request
from apscheduler.schedulers.background import BackgroundScheduler

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
CLAUDE_MODEL = "claude-haiku-4-5-20251001"
DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "radar.db")
SWEEP_INTERVAL_HOURS = int(os.environ.get("SWEEP_INTERVAL_HOURS", "12"))
PORT = int(os.environ.get("PORT", "8080"))

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger("radar")
app = Flask(__name__)
HEADERS = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"}
sweep_running = False
sweep_lock = threading.Lock()


def init_db():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""CREATE TABLE IF NOT EXISTS judgments (
        id TEXT PRIMARY KEY, title TEXT NOT NULL, court TEXT, date_decided TEXT,
        date_fetched TEXT NOT NULL, source TEXT NOT NULL, source_url TEXT,
        full_text_snippet TEXT, categories TEXT, sections TEXT,
        relevance_score TEXT, practitioner_note TEXT, raw_classification TEXT)""")
    conn.execute("""CREATE TABLE IF NOT EXISTS sweep_log (
        id INTEGER PRIMARY KEY AUTOINCREMENT, sweep_time TEXT NOT NULL,
        source TEXT NOT NULL, results_fetched INTEGER, results_relevant INTEGER, errors TEXT)""")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_date_fetched ON judgments(date_fetched)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_categories ON judgments(categories)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_relevance ON judgments(relevance_score)")
    conn.commit()
    conn.close()
    logger.info("Database initialized at %s", DB_PATH)


CLASSIFICATION_PROMPT = """You are a classification engine for Katalyst Advisors, an M&A and transaction structuring advisory firm in India.

IMPORTANT: Be LIBERAL. If there is ANY reasonable connection to M&A, corporate restructuring, family business, corporate governance, taxation of transactions, regulatory action on companies, or deal activity, classify it as relevant. When in doubt, INCLUDE it. Missing something useful is worse than including something borderline.

CATEGORIES:
1. SCHEMES OF ARRANGEMENT: mergers, demergers, amalgamations, composite schemes, capital reductions, NCLT/NCLAT orders on schemes, appointed dates, valuation disputes
2. FAMILY ARRANGEMENT: family settlements, partitions, HUF disputes, succession, family business splits, will contests, inheritance involving business assets
3. INCOME TAX: any tax ruling on M&A, restructuring, capital gains, slump sales, demerger tax, gift tax, deemed income, firm reconstitution, trust taxation, HUF taxation
4. SEBI / SECURITIES: takeover regulations, open offers, inter se transfers, RPTs, insider trading in restructuring context, delisting, promoter reclassification
5. STAMP DUTY: stamp duty on transfers, schemes, conveyances, family arrangements
6. FEMA / CROSS-BORDER: FEMA rulings, cross-border transfers, NRI transactions, FDI in restructuring, RBI directions
7. TRUST LAW: private trusts, charitable trusts in business, trust taxation, settlor-beneficiary disputes
8. INSOLVENCY: IBC matters with promoter families, Section 29A, oppression mismanagement, resolution plans, liquidation
9. BOARDROOM / GOVERNANCE: promoter disputes, board fights, director removals, EGM battles, shareholder activism, SEBI governance enforcement, corporate governance failures
10. IND AS / ACCOUNTING: business combination accounting, Ind AS 103, common control, consolidation, fair value, purchase price allocation, EAC opinions

Respond JSON only (no markdown, no backticks):
{"relevant": true/false, "relevance_score": "high"|"medium"|"low", "categories": [numbers], "category_names": [names], "sections_engaged": [sections if mentioned], "court_or_tribunal": "source or News or Article", "date_decided": "date or null", "parties": "names or null", "summary": "One to two line summary of what this is about."}

If NOT relevant: {"relevant": false, "relevance_score": "none", "reason": "brief reason"}

TEXT:
"""


# ---------------------------------------------------------------------------
# Utilities
# ---------------------------------------------------------------------------

def safe_get(url, timeout=30):
    try:
        resp = requests.get(url, timeout=timeout, headers=HEADERS)
        if resp.status_code == 200:
            return resp
    except Exception as e:
        logger.error("HTTP error for %s: %s", url[:80], str(e)[:100])
    return None

def strip_html(text):
    text = re.sub(r'<[^>]+>', ' ', text)
    return re.sub(r'\s+', ' ', text).strip()

def fetch_google_search(query, num=10):
    """Search Google via Serper.dev API. Returns clean structured results."""
    results = []
    serper_key = os.environ.get("SERPER_API_KEY", "")
    if not serper_key:
        logger.error("No SERPER_API_KEY set. Google search disabled.")
        return results
    try:
        resp = requests.post(
            "https://google.serper.dev/search",
            headers={"X-API-KEY": serper_key, "Content-Type": "application/json"},
            json={"q": query, "num": num, "gl": "in", "hl": "en"},
            timeout=30,
        )
        if resp.status_code == 200:
            data = resp.json()
            for item in data.get("organic", []):
                results.append({
                    "title": item.get("title", ""),
                    "url": item.get("link", ""),
                    "snippet": item.get("snippet", ""),
                    "source": "Google",
                })
            # Also grab news results if present
            for item in data.get("news", []):
                results.append({
                    "title": item.get("title", ""),
                    "url": item.get("link", ""),
                    "snippet": item.get("snippet", ""),
                    "source": "Google News",
                })
        else:
            logger.error("Serper API error %d: %s", resp.status_code, resp.text[:200])
    except Exception as e:
        logger.error("Serper error: %s", str(e)[:100])
    return results

def fetch_rss_feed(feed_url, source_name):
    results = []
    try:
        feed = feedparser.parse(feed_url)
        for entry in feed.entries[:20]:
            results.append({
                "title": entry.get("title", ""),
                "url": entry.get("link", ""),
                "snippet": strip_html(entry.get("summary", entry.get("description", "")))[:1000],
                "source": source_name,
            })
    except Exception as e:
        logger.error("RSS error for %s: %s", source_name, str(e)[:100])
    return results

def scrape_links(url, source_name, min_title_len=15, href_filter=None):
    results = []
    try:
        resp = safe_get(url)
        if resp:
            for href, title_raw in re.findall(r'<a[^>]+href="([^"]*)"[^>]*>(.*?)</a>', resp.text, re.DOTALL):
                title = strip_html(title_raw).strip()
                if len(title) < min_title_len:
                    continue
                if href_filter and not href_filter(href):
                    continue
                full_url = href if href.startswith("http") else urljoin(url, href)
                results.append({"title": title[:200], "url": full_url, "snippet": "", "source": source_name})
    except Exception as e:
        logger.error("Scrape error for %s: %s", source_name, str(e)[:100])
    return results

def fetch_full_text(url):
    if url.lower().endswith(".pdf"):
        return ""
    try:
        resp = safe_get(url)
        if resp:
            ct = resp.headers.get("Content-Type", "")
            if "pdf" in ct or "octet-stream" in ct:
                return ""
            return strip_html(resp.text)[:4000]
    except Exception as e:
        logger.error("Full text error for %s: %s", url[:60], str(e)[:100])
    return ""

def google_multi(queries, source_name, results_per=3, delay=2):
    """Run multiple Google searches and return combined results."""
    results = []
    for q in queries:
        for r in fetch_google_search(q, results_per):
            r["source"] = source_name
            results.append(r)
        time.sleep(delay)
    logger.info("%s: %d results", source_name, len(results))
    return results


# ---------------------------------------------------------------------------
# RSS Feeds
# ---------------------------------------------------------------------------

RSS_FEEDS = [
    ("https://www.livelaw.in/feed", "LiveLaw"),
    ("https://www.barandbench.com/feed", "Bar and Bench"),
    ("https://www.taxmann.com/post/blog/feed/", "Taxmann"),
    ("https://taxguru.in/feed", "TaxGuru"),
    ("https://www.scconline.com/blog/post/feed/", "SCC Online Blog"),
    ("https://www.mondaq.com/india/feeds/rss/latest", "Mondaq India"),
    ("https://corporate.cyrilamarchandblogs.com/feed/", "CAM Blog"),
    ("https://indiacorplaw.in/feed", "IndiaCorpLaw"),
    ("https://taxguru.in/income-tax/feed", "TaxGuru Income Tax"),
    ("https://taxguru.in/company-law/feed", "TaxGuru Company Law"),
    ("https://taxguru.in/sebi/feed", "TaxGuru SEBI"),
    ("https://www.livelaw.in/tax-cases/feed", "LiveLaw Tax"),
    ("https://www.livelaw.in/corporate-law/feed", "LiveLaw Corporate"),
    ("https://www.nishithdesai.com/feed", "Nishith Desai"),
    ("https://erfrequently.com/feed/", "ER Frequently (S&R)"),
    ("https://www.livemint.com/rss/companies", "Livemint Companies"),
    ("https://www.livemint.com/rss/money", "Livemint Money"),
    ("https://economictimes.indiatimes.com/rssfeedstopstories.cms", "Economic Times"),
    ("https://economictimes.indiatimes.com/markets/rssfeeds/1977021501.cms", "ET Markets"),
    ("https://www.business-standard.com/rss/companies-122.rss", "BS Companies"),
    ("https://www.business-standard.com/rss/markets-102.rss", "BS Markets"),
    ("https://www.moneycontrol.com/rss/business.xml", "Moneycontrol"),
    ("https://www.financialexpress.com/feed/", "Financial Express"),
    ("https://www.thehindubusinessline.com/companies/feeder/default.rss", "Hindu BusinessLine"),
    ("https://www.fortuneindia.com/rss/enterprise", "Fortune India"),
    ("https://www.icai.org/rss/rss_feed.html", "ICAI"),
]


# ---------------------------------------------------------------------------
# Source Fetchers (all Google-first except SEBI which scrapes well)
# ---------------------------------------------------------------------------

def fetch_indian_kanoon(query):
    results = []
    try:
        api_key = os.environ.get("INDIAN_KANOON_API_KEY", "")
        if api_key:
            params = {"formInput": query, "pagenum": 0}
            resp = requests.get("https://api.indiankanoon.org/search/", params=params,
                                headers={"Authorization": "Token " + api_key}, timeout=30)
            if resp.status_code == 200:
                for doc in resp.json().get("docs", []):
                    results.append({"title": doc.get("title", ""),
                                    "url": "https://indiankanoon.org/doc/" + str(doc.get("tid", "")) + "/",
                                    "snippet": doc.get("headline", ""), "source": "Indian Kanoon"})
        else:
            for r in fetch_google_search("site:indiankanoon.org " + query, 5):
                if "indiankanoon.org" in r["url"]:
                    r["source"] = "Indian Kanoon"
                    results.append(r)
    except Exception as e:
        logger.error("IK error: %s", str(e)[:100])
    return results


def fetch_nclt():
    return google_multi([
        "NCLT order scheme arrangement 2026",
        "NCLT order demerger 2026",
        "NCLT order amalgamation 2026",
        "NCLT order capital reduction 2026",
        "NCLT order oppression mismanagement 2026",
        "NCLT order company petition 2026",
        "NCLT Mumbai bench order 2026",
        "NCLT Delhi bench order 2026",
        "NCLT Ahmedabad order 2026",
        "NCLT Bengaluru order 2026",
        "NCLT Chennai order 2026",
        "NCLT Kolkata order 2026",
        "NCLT Hyderabad order 2026",
        "NCLT Chandigarh order 2026",
        "NCLT scheme sanction order India",
        "NCLT appointed date scheme order",
        "NCLT winding up order 2026",
        "site:livelaw.in NCLT order 2026",
        "site:barandbench.com NCLT order 2026",
        "site:scconline.com NCLT order 2026",
        "site:taxguru.in NCLT order 2026",
    ], "NCLT", results_per=5, delay=2)


def fetch_nclat():
    return google_multi([
        "NCLAT order 2026",
        "NCLAT appeal scheme arrangement 2026",
        "NCLAT judgment company law 2026",
        "NCLAT order insolvency 2026",
        "NCLAT order oppression mismanagement",
        "site:livelaw.in NCLAT 2026",
        "site:barandbench.com NCLAT 2026",
        "site:scconline.com NCLAT 2026",
    ], "NCLAT", results_per=5, delay=2)


def fetch_supreme_court():
    return google_multi([
        "Supreme Court India company law judgment 2026",
        "Supreme Court India merger demerger 2026",
        "Supreme Court India family settlement 2026",
        "Supreme Court India corporate dispute 2026",
        "Supreme Court India tax appeal restructuring 2026",
        "Supreme Court India SEBI appeal 2026",
        "Supreme Court India insolvency IBC 2026",
        "Supreme Court India capital gains exemption 2026",
        "site:main.sci.gov.in judgment 2026",
        "site:livelaw.in Supreme Court company law 2026",
        "site:barandbench.com Supreme Court corporate 2026",
        "site:scconline.com Supreme Court company 2026",
    ], "Supreme Court", results_per=5, delay=2)


def fetch_high_courts():
    return google_multi([
        "Bombay High Court scheme arrangement 2026",
        "Bombay High Court company petition 2026",
        "Bombay High Court stamp duty scheme",
        "Bombay High Court family settlement",
        "Bombay High Court writ NCLT",
        "Delhi High Court scheme arrangement 2026",
        "Delhi High Court family arrangement 2026",
        "Delhi High Court company law 2026",
        "Gujarat High Court scheme arrangement 2026",
        "Gujarat High Court family settlement stamp duty",
        "Gujarat High Court company petition 2026",
        "Karnataka High Court scheme arrangement 2026",
        "Madras High Court company petition 2026",
        "Madras High Court scheme arrangement 2026",
        "Calcutta High Court company petition 2026",
        "Allahabad High Court family arrangement",
        "Punjab Haryana High Court company petition 2026",
        "Kerala High Court scheme arrangement 2026",
        "Telangana High Court company petition 2026",
        "High Court writ petition NCLT order 2026",
        "High Court appeal NCLT scheme 2026",
        "High Court capital gains merger exemption India",
        "High Court stamp duty merger exemption India",
        "High Court corporate dispute India 2026",
        "High Court family partition property India 2026",
    ], "High Court", results_per=3, delay=2)


def fetch_sebi():
    """SEBI website works with direct scraping."""
    results = []
    for sebi_url, name in [
        ("https://www.sebi.gov.in/sebiweb/home/HomeAction.do?doListing=yes&sid=1&ssid=2&smid=0", "SEBI Orders"),
        ("https://www.sebi.gov.in/sebiweb/home/HomeAction.do?doListing=yes&sid=2&ssid=9&smid=0", "SEBI Circulars"),
    ]:
        results.extend(scrape_links(sebi_url, name, min_title_len=15))
        time.sleep(1)
    # Supplement with Google
    results.extend(google_multi([
        "SEBI order 2026 corporate governance",
        "SEBI order 2026 takeover",
        "SEBI order 2026 insider trading",
        "SEBI circular 2026 listing",
        "SEBI order related party transaction 2026",
        "SEBI adjudication order 2026",
        "site:sebi.gov.in order 2026",
    ], "SEBI (Google)", results_per=3, delay=2))
    logger.info("SEBI total: %d", len(results))
    return results


def fetch_sat():
    return google_multi([
        "SAT order 2026",
        "Securities Appellate Tribunal order 2026",
        "SAT SEBI appeal 2026",
        "SAT order takeover open offer",
        "SAT order insider trading",
        "site:sat.gov.in order 2026",
        "site:livelaw.in SAT order 2026",
    ], "SAT", results_per=3, delay=2)


def fetch_itat():
    return google_multi([
        "ITAT order demerger section 2(19AA)",
        "ITAT order slump sale section 50B",
        "ITAT order family partition HUF capital gains",
        "ITAT order trust taxation section 161",
        "ITAT order capital gains section 47 exemption",
        "ITAT order deemed gift section 56(2)(x)",
        "ITAT order reconstitution section 45(4)",
        "ITAT order dissolution section 9B",
        "ITAT order amalgamation tax",
        "ITAT order transfer pricing restructuring",
        "site:itatonline.org 2026",
        "site:taxguru.in ITAT order 2026",
        "site:livelaw.in ITAT order 2026",
    ], "ITAT", results_per=3, delay=1.5)


def fetch_cci():
    return google_multi([
        "CCI combination approval order 2026",
        "Competition Commission merger India 2026",
        "CCI order combination review India",
        "site:cci.gov.in combination order",
    ], "CCI", results_per=5, delay=1.5)


def fetch_mca_ibbi():
    return google_multi([
        "MCA notification Companies Act 2026",
        "MCA circular company law 2026",
        "IBBI circular regulation 2026",
        "IBBI valuation standard 2026",
        "site:mca.gov.in notification 2026",
        "site:ibbi.gov.in circular 2026",
    ], "MCA/IBBI", results_per=3, delay=1.5)


def fetch_exchange():
    return google_multi([
        # Search for news coverage of exchange filings, not the exchange sites directly
        "BSE filing scheme of arrangement India 2026",
        "NSE filing scheme of arrangement India 2026",
        "BSE filing demerger India 2026",
        "stock exchange scheme arrangement filing India",
        "company filed scheme arrangement NCLT 2026",
        "open offer letter filed SEBI 2026",
        "delisting offer filed stock exchange India 2026",
        "composite scheme filed BSE NSE 2026",
        "corporate action India demerger merger 2026",
        "scheme of arrangement SEBI no objection 2026",
        "SEBI observation letter scheme arrangement 2026",
        "listed company demerger announcement India 2026",
        "listed company merger announcement India 2026",
        "listed company open offer announcement India 2026",
        "promoter open offer India 2026",
        "voluntary delisting India 2026",
        "record date demerger India 2026",
        "entitlement ratio demerger India 2026",
    ], "Exchange / Deal Filings", results_per=3, delay=2)


def fetch_newspaper_governance():
    return google_multi([
        "site:livemint.com boardroom battle 2026",
        "site:livemint.com promoter family dispute 2026",
        "site:livemint.com demerger scheme 2026",
        "site:livemint.com corporate governance 2026",
        "site:livemint.com company restructuring 2026",
        "site:livemint.com merger acquisition 2026",
        "site:livemint.com family business succession 2026",
        "site:economictimes.indiatimes.com promoter feud 2026",
        "site:economictimes.indiatimes.com boardroom coup 2026",
        "site:economictimes.indiatimes.com family business split 2026",
        "site:economictimes.indiatimes.com M&A deal India 2026",
        "site:economictimes.indiatimes.com corporate restructuring 2026",
        "site:business-standard.com family business dispute 2026",
        "site:business-standard.com promoter restructuring 2026",
        "site:business-standard.com merger acquisition India 2026",
        "site:business-standard.com corporate governance 2026",
        "site:moneycontrol.com boardroom battle 2026",
        "site:moneycontrol.com promoter dispute 2026",
        "site:moneycontrol.com merger demerger India 2026",
        "site:moneycontrol.com corporate restructuring 2026",
    ], "Business News", results_per=3, delay=2)


def fetch_law_firm_blogs():
    results = []
    for firm in ["site:azbpartners.com", "site:khaitanco.com", "site:trilegal.com",
                  "site:nishithdesai.com", "site:cyrilamarchandblogs.com"]:
        for topic in ["scheme arrangement", "family settlement", "SEBI takeover",
                       "M&A India", "corporate restructuring", "tax restructuring",
                       "demerger", "merger", "open offer"]:
            for r in fetch_google_search(firm + " " + topic, 2):
                r["source"] = "Law Firm Blog"
                results.append(r)
            time.sleep(1.5)
    logger.info("Law firms: %d", len(results))
    return results


def fetch_proxy_advisory():
    return google_multi([
        # Search for news coverage of proxy advisory actions, not the paywalled reports
        "IiAS opposes related party transaction India",
        "IiAS recommends against resolution India",
        "InGovern raises governance concerns India",
        "SES Governance proxy advisory India",
        "proxy advisory opposes merger India",
        "proxy advisory opposes scheme India",
        "proxy advisory flags related party India",
        "proxy advisory firm India 2026 recommendation",
        "institutional investor opposes resolution India 2026",
        "IiAS voting advisory India 2026",
        "proxy advisory delisting India",
        "proxy advisory open offer India",
        "minority shareholders oppose scheme India 2026",
        "institutional shareholder concerns India governance",
    ], "Proxy Advisory", results_per=3, delay=1.5)


def fetch_indas():
    return google_multi([
        "Ind AS 103 business combination India",
        "common control transaction accounting India",
        "purchase price allocation India Ind AS",
        "deferred tax business combination India",
        "ICAI EAC opinion scheme arrangement",
        "goodwill impairment India Ind AS",
        "opening balance sheet demerger accounting",
        "accounting amalgamation India Ind AS",
        "fair value measurement Ind AS 113",
        "site:deloitte.com/in Ind AS",
        "site:pwc.in Ind AS accounting",
        "site:ey.com/en_in Ind AS",
        "site:kpmg.com/in Ind AS",
    ], "Ind AS / Accounting", results_per=3, delay=1.5)


# ---------------------------------------------------------------------------
# Search Query Lists
# ---------------------------------------------------------------------------

SEARCH_QUERIES_TECHNICAL = [
    "section 230 Companies Act scheme arrangement",
    "section 232 Companies Act amalgamation",
    "section 66 capital reduction NCLT",
    "section 2(19AA) demerger income tax",
    "section 47 exemption amalgamation transfer",
    "section 56(2)(x) deemed gift",
    "section 45(4) reconstitution firm",
    "section 9B dissolution partnership LLP",
    "section 50B slump sale",
    "section 50D fair market value",
    "section 171 HUF partition",
    "section 47(iii) gift transfer",
    "SEBI regulation 10 inter se transfer promoter",
    "SEBI regulation 23 related party transaction",
    "SEBI regulation 37 scheme listing",
    "SEBI takeover code open offer regulation 3",
    "SEBI delisting regulation",
    "section 29A IBC related party resolution",
    "section 241 oppression mismanagement Companies Act",
    "Ind AS 103 business combination",
    "Ind AS 110 consolidation",
    "Ind AS 12 deferred tax restructuring",
    "Ind AS 113 fair value measurement",
    "Ind AS 27 separate financial statements",
]

SEARCH_QUERIES_PLAIN_ENGLISH = [
    "company merger India 2026",
    "company demerger India 2026",
    "corporate restructuring India 2026",
    "business restructuring India tax",
    "group restructuring India promoter",
    "family business split India",
    "family business dispute India",
    "family business succession India",
    "promoter family fight India",
    "promoter shareholding change India",
    "promoter group restructuring India",
    "listed company restructuring India",
    "business transfer agreement India",
    "slump sale India deal",
    "share swap merger India",
    "holding subsidiary restructuring India",
    "private equity buyout India",
    "open offer India 2026",
    "delisting India promoter",
    "corporate governance India 2026",
    "boardroom fight India",
    "independent director controversy India",
    "shareholder dispute India",
    "minority shareholder rights India",
    "related party transaction abuse India",
    "business valuation dispute India",
    "fair value dispute company India",
    "capital gains exemption merger India",
    "tax free merger India",
    "succession planning wealthy family India",
    "insolvency promoter family India",
    "IBC resolution plan India 2026",
    "NRI property transfer FEMA India",
    "family trust India tax",
    "wealth planning India trust",
    "stamp duty merger India 2026",
    "composite scheme India 2026",
    "appointed date scheme India",
    "company acquisition India 2026",
    "takeover India 2026",
]

SEARCH_QUERIES_GOVERNANCE = [
    "boardroom battle India 2026",
    "promoter family feud listed company India",
    "corporate governance failure SEBI India",
    "SEBI penalty governance India 2026",
    "independent director removal India 2026",
    "shareholder activism India 2026",
    "proxy advisory India governance 2026",
    "promoter reclassification SEBI India",
    "EGM requisition India promoter dispute",
    "board coup India company 2026",
    "family succession corporate dispute India",
    "pledge enforcement promoter shares India",
]


# ---------------------------------------------------------------------------
# Classification and Storage
# ---------------------------------------------------------------------------

def classify_judgment(text, title=""):
    if not ANTHROPIC_API_KEY:
        logger.error("No ANTHROPIC_API_KEY set.")
        return None
    combined = ("TITLE: " + title + "\n\nTEXT:\n" + text) if title else text
    try:
        resp = requests.post("https://api.anthropic.com/v1/messages",
            headers={"Content-Type": "application/json", "x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01"},
            json={"model": CLAUDE_MODEL, "max_tokens": 512, "messages": [{"role": "user", "content": CLASSIFICATION_PROMPT + combined}]},
            timeout=60)
        if resp.status_code == 200:
            content = resp.json().get("content", [{}])[0].get("text", "").strip().strip("`").strip()
            if content.startswith("json"):
                content = content[4:].strip()
            return json.loads(content)
        else:
            logger.error("Claude API error %d: %s", resp.status_code, resp.text[:200])
    except json.JSONDecodeError as e:
        logger.error("JSON parse error: %s", e)
    except Exception as e:
        logger.error("Classification error: %s", e)
    return None

def generate_id(title, url):
    return hashlib.sha256((title + "|" + url).encode()).hexdigest()[:16]

def is_already_processed(jid):
    conn = sqlite3.connect(DB_PATH)
    exists = conn.execute("SELECT 1 FROM judgments WHERE id = ?", (jid,)).fetchone() is not None
    conn.close()
    return exists

def save_judgment(jid, raw, cls):
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""INSERT OR REPLACE INTO judgments
        (id,title,court,date_decided,date_fetched,source,source_url,full_text_snippet,
         categories,sections,relevance_score,practitioner_note,raw_classification)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (jid, raw.get("title","Unknown"), cls.get("court_or_tribunal",""),
         cls.get("date_decided",""), datetime.utcnow().isoformat(),
         raw.get("source",""), raw.get("url",""), raw.get("snippet","")[:500],
         json.dumps(cls.get("categories",[]),separators=(",",":")), json.dumps(cls.get("sections_engaged",[]),separators=(",",":")),
         cls.get("relevance_score",""), cls.get("summary", cls.get("practitioner_note","")),
         json.dumps(cls)))
    conn.commit()
    conn.close()

def log_sweep(source, fetched, relevant, errors=""):
    conn = sqlite3.connect(DB_PATH)
    conn.execute("INSERT INTO sweep_log (sweep_time,source,results_fetched,results_relevant,errors) VALUES (?,?,?,?,?)",
        (datetime.utcnow().isoformat(), source, fetched, relevant, errors))
    conn.commit()
    conn.close()

def process_results(results, label):
    count = 0
    seen = set()
    unique = [r for r in results if r["url"] not in seen and not seen.add(r["url"])]
    logger.info("%s: %d unique to process", label, len(unique))
    for result in unique:
        jid = generate_id(result["title"], result["url"])
        if is_already_processed(jid):
            continue
        text = result.get("snippet", "")
        if not text or len(text) < 80:
            text = fetch_full_text(result["url"])
            time.sleep(0.5)
        if not text or len(text) < 30:
            continue
        cls = classify_judgment(text, result["title"])
        if cls and cls.get("relevant"):
            save_judgment(jid, result, cls)
            count += 1
            logger.info("  RELEVANT: %s [%s]", result["title"][:80], cls.get("relevance_score",""))
        time.sleep(0.3)
    log_sweep(label, len(unique), count)
    return count


# ---------------------------------------------------------------------------
# Sweep
# ---------------------------------------------------------------------------

def run_sweep():
    global sweep_running
    acquired = sweep_lock.acquire(blocking=False)
    if not acquired:
        logger.info("Sweep already running, skipping.")
        return
    sweep_running = True
    try:
        _do_sweep()
    finally:
        sweep_running = False
        sweep_lock.release()

def _do_sweep():
    logger.info("=== v5-FINAL SWEEP at %s ===", datetime.utcnow().isoformat())
    total = 0

    logger.info("--- Phase 1: RSS Feeds (26 feeds) ---")
    rss = []
    for feed_url, name in RSS_FEEDS:
        logger.info("  %s...", name)
        rss.extend(fetch_rss_feed(feed_url, name))
        time.sleep(0.5)
    total += process_results(rss, "RSS Feeds")

    logger.info("--- Phase 2: Indian Kanoon (technical queries) ---")
    ik = []
    for q in SEARCH_QUERIES_TECHNICAL:
        ik.extend(fetch_indian_kanoon(q))
        time.sleep(1.5)
    total += process_results(ik, "Indian Kanoon")

    logger.info("--- Phase 3: Courts via Google ---")
    total += process_results(fetch_nclt(), "NCLT")
    total += process_results(fetch_nclat(), "NCLAT")
    total += process_results(fetch_supreme_court(), "Supreme Court")
    total += process_results(fetch_high_courts(), "High Courts")

    logger.info("--- Phase 4: Regulators ---")
    total += process_results(fetch_sebi(), "SEBI")
    total += process_results(fetch_sat(), "SAT")
    total += process_results(fetch_itat(), "ITAT")
    total += process_results(fetch_cci(), "CCI")
    total += process_results(fetch_mca_ibbi(), "MCA/IBBI")
    total += process_results(fetch_exchange(), "Exchanges")

    logger.info("--- Phase 5: Plain English Google (40 queries) ---")
    pe = []
    for q in SEARCH_QUERIES_PLAIN_ENGLISH:
        pe.extend(fetch_google_search(q, 3))
        time.sleep(2)
    total += process_results(pe, "Plain English")

    logger.info("--- Phase 6: Ind AS / Accounting ---")
    total += process_results(fetch_indas(), "Ind AS")

    logger.info("--- Phase 7: Governance ---")
    gov = []
    for q in SEARCH_QUERIES_GOVERNANCE:
        gov.extend(fetch_google_search(q, 3))
        time.sleep(2)
    total += process_results(gov, "Governance")
    total += process_results(fetch_newspaper_governance(), "Newspaper Governance")
    total += process_results(fetch_proxy_advisory(), "Proxy Advisory")

    logger.info("--- Phase 8: Law Firm Blogs ---")
    total += process_results(fetch_law_firm_blogs(), "Law Firm Blogs")

    logger.info("=== SWEEP COMPLETE. Total relevant: %d ===", total)


# ---------------------------------------------------------------------------
# API
# ---------------------------------------------------------------------------

@app.route("/")
def index():
    p = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static", "index.html")
    with open(p, "r") as f:
        return f.read(), 200, {"Content-Type": "text/html; charset=utf-8"}

@app.route("/api/judgments")
def get_judgments():
    category = request.args.get("category", "")
    score = request.args.get("score", "")
    days = int(request.args.get("days", "365"))
    search = request.args.get("search", "")
    page = int(request.args.get("page", "1"))
    per_page = 20
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    wc, params = [], []
    wc.append("date_fetched >= ?")
    params.append((datetime.utcnow() - timedelta(days=days)).isoformat())
    if category:
        wc.append("(categories LIKE ? OR categories LIKE ? OR categories LIKE ? OR categories LIKE ? OR categories LIKE ? OR categories LIKE ?)")
        params.extend([
            "[" + category + "]",
            "[" + category + ",",
            "," + category + ",",
            "," + category + "]",
            ", " + category + ",",
            ", " + category + "]",
        ])
    if score:
        wc.append("relevance_score = ?"); params.append(score)
    if search:
        wc.append("(title LIKE ? OR practitioner_note LIKE ? OR sections LIKE ?)")
        params.extend(["%" + search + "%"] * 3)
    w = " AND ".join(wc) if wc else "1=1"
    offset = (page - 1) * per_page
    total = conn.execute("SELECT COUNT(*) as cnt FROM judgments WHERE " + w, params).fetchone()["cnt"]
    rows = conn.execute("SELECT * FROM judgments WHERE " + w + " ORDER BY date_fetched DESC LIMIT ? OFFSET ?", params + [per_page, offset]).fetchall()
    conn.close()
    jlist = []
    for row in rows:
        jlist.append({"id": row["id"], "title": row["title"], "court": row["court"],
            "date_decided": row["date_decided"], "date_fetched": row["date_fetched"],
            "source": row["source"], "source_url": row["source_url"],
            "categories": json.loads(row["categories"]) if row["categories"] else [],
            "sections": json.loads(row["sections"]) if row["sections"] else [],
            "relevance_score": row["relevance_score"], "practitioner_note": row["practitioner_note"]})
    return jsonify({"judgments": jlist, "total": total, "page": page, "per_page": per_page, "total_pages": (total + per_page - 1) // per_page})

@app.route("/api/stats")
def get_stats():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    total = conn.execute("SELECT COUNT(*) as cnt FROM judgments").fetchone()["cnt"]
    high = conn.execute("SELECT COUNT(*) as cnt FROM judgments WHERE relevance_score = 'high'").fetchone()["cnt"]
    ls = conn.execute("SELECT sweep_time FROM sweep_log ORDER BY id DESC LIMIT 1").fetchone()
    cats = {}
    for row in conn.execute("SELECT categories FROM judgments").fetchall():
        for c in (json.loads(row["categories"]) if row["categories"] else []):
            cats[str(c)] = cats.get(str(c), 0) + 1
    conn.close()
    return jsonify({"total_judgments": total, "high_relevance": high, "last_sweep": ls["sweep_time"] if ls else None, "category_distribution": cats})

@app.route("/api/export")
def export_doc():
    """Export current filtered results as a Word document."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    category = request.args.get("category", "")
    score = request.args.get("score", "")
    days = int(request.args.get("days", "365"))
    search = request.args.get("search", "")

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    wc, params = [], []
    wc.append("date_fetched >= ?")
    params.append((datetime.utcnow() - timedelta(days=days)).isoformat())
    if category:
        wc.append("(categories LIKE ? OR categories LIKE ? OR categories LIKE ? OR categories LIKE ? OR categories LIKE ? OR categories LIKE ?)")
        params.extend(["[" + category + "]", "[" + category + ",", "," + category + ",", "," + category + "]", ", " + category + ",", ", " + category + "]"])
    if score:
        wc.append("relevance_score = ?"); params.append(score)
    if search:
        wc.append("(title LIKE ? OR practitioner_note LIKE ? OR sections LIKE ?)")
        params.extend(["%" + search + "%"] * 3)
    w = " AND ".join(wc) if wc else "1=1"
    rows = conn.execute("SELECT * FROM judgments WHERE " + w + " ORDER BY date_fetched DESC LIMIT 200", params).fetchall()
    conn.close()

    cat_names = {1:"Schemes",2:"Family Arrangement",3:"Income Tax",4:"SEBI",5:"Stamp Duty",6:"FEMA",7:"Trust Law",8:"Insolvency",9:"Boardroom",10:"Ind AS"}

    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Katalyst Judgment Radar")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(139, 37, 0)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Generated: " + datetime.utcnow().strftime("%d %B %Y") + " | Results: " + str(len(rows)))
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(107, 101, 96)

    doc.add_paragraph("")

    for row in rows:
        # Title
        p = doc.add_paragraph()
        t_run = p.add_run(row["title"] or "Untitled")
        t_run.bold = True
        t_run.font.size = Pt(11)
        t_run.font.color.rgb = RGBColor(26, 26, 26)

        # Meta line
        meta_parts = []
        if row["court"]:
            meta_parts.append(row["court"])
        if row["source"]:
            meta_parts.append(row["source"])
        if row["date_fetched"]:
            try:
                dt = datetime.fromisoformat(row["date_fetched"])
                meta_parts.append(dt.strftime("%d %b %Y"))
            except:
                pass
        relevance = (row["relevance_score"] or "").upper()
        if relevance:
            meta_parts.append("Relevance: " + relevance)
        if meta_parts:
            mp = doc.add_paragraph()
            mr = mp.add_run(" | ".join(meta_parts))
            mr.font.size = Pt(9)
            mr.font.color.rgb = RGBColor(107, 101, 96)

        # Categories and sections
        tags = []
        try:
            for c in json.loads(row["categories"] or "[]"):
                tags.append(cat_names.get(c, "Cat " + str(c)))
        except:
            pass
        try:
            for s in json.loads(row["sections"] or "[]"):
                tags.append(str(s))
        except:
            pass
        if tags:
            tp = doc.add_paragraph()
            tr = tp.add_run(", ".join(tags))
            tr.font.size = Pt(9)
            tr.italic = True
            tr.font.color.rgb = RGBColor(71, 85, 105)

        # Summary
        if row["practitioner_note"]:
            sp = doc.add_paragraph()
            sr = sp.add_run(row["practitioner_note"])
            sr.font.size = Pt(10)

        # Link
        if row["source_url"]:
            lp = doc.add_paragraph()
            lr = lp.add_run(row["source_url"])
            lr.font.size = Pt(8)
            lr.font.color.rgb = RGBColor(139, 37, 0)

        doc.add_paragraph("_" * 60)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    from flask import send_file
    return send_file(buf, as_attachment=True,
                     download_name="Katalyst_Radar_" + datetime.utcnow().strftime("%Y%m%d") + ".docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/api/sweep", methods=["POST"])
def trigger_sweep():
    if sweep_lock.locked():
        return jsonify({"status": "Sweep already running. Please wait."}), 409
    threading.Thread(target=run_sweep).start()
    return jsonify({"status": "Sweep started", "time": datetime.utcnow().isoformat()})

init_db()
scheduler = BackgroundScheduler()
scheduler.add_job(run_sweep, "interval", hours=SWEEP_INTERVAL_HOURS, id="main_sweep",
                  next_run_time=datetime.utcnow() + timedelta(minutes=2))
scheduler.start()
logger.info("v5-final started. Sweeps every %d hours.", SWEEP_INTERVAL_HOURS)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=PORT, debug=False)
